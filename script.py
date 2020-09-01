from flask import Flask, render_template, send_file, request
from docx import Document
from docxtpl import DocxTemplate, InlineImage
import io
from io import BytesIO
import os
from sap.cf_logging import flask_logging
import logging
from flask_cors import CORS
import json
from werkzeug.exceptions import HTTPException

app = Flask(__name__)
flask_logging.init(app, logging.INFO)
cors = CORS(app, resources={r"/api/*": {"origins": "*"}})
cors = CORS(app, resources={r"/genrate_temp/*": {"origins": "*"}})
@app.route('/api/',methods = ['POST'])
def index():
    data = request

    context = request.json# gets the context used to render the document
    template = context['root']['tem'] # template that will be used
    ##
    cont= context['root']['processList']['processList']
    
        
    for x in cont:   
        x1 = x['callActivity']
        for y in x1: 
            for key, value in y.items():   
                if key == "Content Modifier":
                    if 'Properties' in value:
                        value['Properties'] = json.loads(value['Properties'])
                    if 'Headers' in value:   
                        value['Headers'] = json.loads(value['Headers'])
                    
    
    for x in cont:   
        x1 = x['subProcess']
        for y in x1: 
            for key, value in y.items():   
                if key == "Content Modifier":
                    if 'Properties' in value:
                        value['Properties'] = json.loads(value['Properties'])
                    if 'Headers' in value:   
                        value['Headers'] = json.loads(value['Headers'])


    
    # saving tamplet localy 
    #source_stream = io.BytesIO(bytes.fromhex(template))
    #document = Document(source_stream)
    #source_stream.close()
    #document.save("primaIncercare.docx")

    source_stream = io.BytesIO(bytes.fromhex(template))
    ##
    for key in context['root']['resource']['HM_resource'].keys():
        for res in context['root']['resource']['HM_resource'][key]:
            res['resource'] = bytes.fromhex(res['resource']).decode('utf-8')
            
          
        
        #res.resource = io.BytesIO(bytes.fromhex(res.resource))
    
    embedded_docx_tpl=DocxTemplate('embedded_embedded_docx_map_tpl.docx')
    embedded_docx_tpl.render(context,autoescape=True)
    embedded_docx_tpl.save('embedded_embedded_map.docx')

    embedded_docx_tpl=DocxTemplate('embedded_embedded_docx_groovy_tpl.docx')
    embedded_docx_tpl.render(context,autoescape=True)
    embedded_docx_tpl.save('embedded_embedded_groovy.docx')

    embedded_docx_tpl=DocxTemplate('embedded_embedded_docx_wsdl_tpl.docx')
    embedded_docx_tpl.render(context,autoescape=True)
    embedded_docx_tpl.save('embedded_embedded_wsdl.docx')

    embedded_docx_tpl=DocxTemplate('embedded_embedded_docx_xsd_tpl.docx')
    embedded_docx_tpl.render(context,autoescape=True)
    embedded_docx_tpl.save('embedded_embedded_xsd.docx')

    template = DocxTemplate('MyTemplate.docx')
    #template = DocxTemplate(source_stream)

    target_file = io.BytesIO()
    template.replace_embedded('embedded_dummy_map.docx','embedded_embedded_map.docx')
    template.replace_embedded('embedded_dummy_groovy.docx','embedded_embedded_groovy.docx')
    template.replace_embedded('embedded_dummy_wsdl.docx','embedded_embedded_wsdl.docx')
    template.replace_embedded('embedded_dummy_xsd.docx','embedded_embedded_xsd.docx')
    template.render(context,autoescape=True)
    template.save(target_file)
    
    
    target_file.seek(0)
   
    return send_file(target_file, as_attachment=True, attachment_filename='report.docx')
    

@app.route('/genrate_temp/',methods = ['GET'])
def index1():
    data = request
    param = request.args
    #context = request.json# gets the context used to render the document

    target_file = io.BytesIO()
    document = Document()
     
    for i in param.keys():
        if(param[i] == 'yes'):
            path = str(i) + ".docx"
            source_document = Document(path)
            for element in source_document.element.body:
                document.element.body.append(element)
    
    
    document.save(target_file)
    target_file.seek(0)
    return send_file(target_file, as_attachment=True, attachment_filename='report.docx')

@app.errorhandler(HTTPException)
def handle_exception(e):
    """Return JSON instead of HTML for HTTP errors."""
    # start with the correct headers and status code from the error
    response = e.get_response()
    # replace the body with JSON
    response.data = json.dumps({
        "code": e.code,
        "name": e.name,
        "description": "Something goes wrong on the server-side. Please revision the template configurations.",
    })
    response.content_type = "application/json"
    response.status = '200'
    return response


#if __name__ == "__main__":
#    app.run(debug=True)

if __name__ == "__main__":
    osPort = os.getenv("PORT",)
    if osPort == None:
        port = 5000
    else:
        port = int(osPort)
    app.run(host='0.0.0.0', port=port)